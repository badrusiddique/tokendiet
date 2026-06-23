"""Evidence that Markdown is the right target format.

For each input type we count tokens (tiktoken cl100k proxy) for every candidate
encoding, plus the "native" cost Claude pays without conversion. The takeaways:

* base64 / binary are 7x-360x WORSE and unreadable — never encode bytes as text.
* Markdown ≈ plain text on prose; Markdown wins by also keeping links/structure.
* The real lever is the per-input extractor (all targeting Markdown), not the
  container: image->OCR and (analytically) video->transcript are the big wins.

Run:  python benchmarks/format_compare.py
(DOCX row needs the probe group + markitdown; it is skipped if unavailable.)
"""

from __future__ import annotations

import base64
import logging
import tempfile
from pathlib import Path

import pymupdf
import tiktoken

from tokendiet.html import html_to_markdown
from tokendiet.tokens import image_tokens, page_image_tokens

logging.disable(logging.CRITICAL)
enc = tiktoken.get_encoding("cl100k_base")
CORPUS = Path(__file__).parent / "corpus"


def t(s: str) -> int:
    return len(enc.encode(s))


def b64(raw: bytes) -> int:
    return t(base64.standard_b64encode(raw).decode())


def show(label: str, tokens: int, base: int, note: str = "") -> None:
    rel = f"{tokens / base:.1f}x" if base else "-"
    print(f"   {label:26} {tokens:>10,} tok  {rel:>7} md  {note}")


def pdf_section(path: Path) -> None:
    doc = pymupdf.open(path)
    import pymupdf4llm

    md = pymupdf4llm.to_markdown(doc, show_progress=False)
    plain = "\n".join(p.get_text() for p in doc)
    native = sum(page_image_tokens(p.rect.width, p.rect.height) for p in doc) + t(md)
    raw = path.read_bytes()
    doc.close()
    base = t(md)
    print(f"\nPDF: {path.name}")
    show("native (text+images)", native, base, "← without Tokendiet")
    show("markdown", base, base, "← default")
    show("plain text", t(plain), base)
    show("base64 (bytes)", b64(raw), base, "✗ unreadable")


def html_section(path: Path) -> None:
    import trafilatura

    raw = path.read_text(encoding="utf-8", errors="replace")
    md = html_to_markdown(raw)
    txt = trafilatura.extract(raw, output_format="txt") or ""
    base = t(md)
    print(f"\nHTML: {path.name}")
    show("raw HTML (native)", t(raw), base, "← without Tokendiet")
    show("markdown (main-content)", base, base, "← default")
    show("plain text", t(txt), base, "cheaper, but drops links")
    show("base64 (bytes)", b64(raw.encode()), base, "✗")


def image_section() -> None:
    pdf = CORPUS / "prose.pdf"
    if not pdf.exists():
        return
    doc = pymupdf.open(pdf)
    page = doc[0]
    pix = page.get_pixmap(dpi=150)
    with tempfile.TemporaryDirectory() as tmp:
        png = Path(tmp) / "p.png"
        pix.save(png)
        raw = png.read_bytes()
    ocr_proxy = page.get_text()  # perfect-OCR proxy = the page's real text
    doc.close()
    base = t(ocr_proxy)
    print(f"\nIMAGE: a page rendered to PNG {pix.width}x{pix.height}")
    show("send as image (native)", image_tokens(pix.width, pix.height), base, "← without Tokendiet")
    show("OCR -> text/markdown", base, base, "← default (if only text needed)")
    show("base64 (png bytes)", b64(raw), base, "✗")


def docx_section() -> None:
    try:
        from docx import Document
        from markitdown import MarkItDown
    except ImportError:
        print("\nDOCX: skipped (needs `uv sync --group probe` + markitdown)")
        return
    prose = "Tokendiet keeps the body and drops the chrome. " * 30
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "s.docx"
        d = Document()
        for _ in range(15):
            d.add_paragraph(prose)
        d.save(p)
        md = MarkItDown().convert(str(p)).text_content
        plain = "\n".join(par.text for par in Document(p).paragraphs)
        raw = p.read_bytes()
    base = t(md)
    print("\nDOCX: generated")
    show("markdown", base, base)
    show("plain text", t(plain), base, "≈ markdown (no native image/markup cost)")
    show("base64 (bytes)", b64(raw), base, "✗")


def video_section() -> None:
    print("\nVIDEO: 3-min clip (analytical — no transcription engine bundled)")
    frames, per_frame = 180, image_tokens(1568, 882)
    transcript = int(450 * 1.3)
    show("send frames (native)", frames * per_frame, transcript, f"← {frames} frames as images")
    show("transcript (text/md)", transcript, transcript, "← ASR; ~567x cheaper")


def main() -> int:
    for name in ("prose.pdf", "paper-attention.pdf"):
        if (CORPUS / name).exists():
            pdf_section(CORPUS / name)
    for name in ("wikipedia-markdown.html", "article.html"):
        if (CORPUS / name).exists():
            html_section(CORPUS / name)
    image_section()
    docx_section()
    video_section()
    print(
        "\nVerdict: Markdown is the right container. base64/binary are far worse. "
        "Per-input extraction (OCR, transcript) is where the wins are."
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
